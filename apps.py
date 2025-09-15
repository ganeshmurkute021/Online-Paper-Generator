import streamlit as st
import pandas as pd
import random
from docx import Document
from io import BytesIO
from docx.shared import Inches, Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Load dataset
try:
    questions_df = pd.read_csv(r"C:\Users\Anju\Downloads\questions\question_dataset.csv")  # Replace with your dataset path
except FileNotFoundError:
    st.error("Error: File not found. Please ensure the file exists in the specified path.")
    st.stop()


# Generate questions for a section
def generate_questions(filtered_questions, num_questions, marks):
    available_questions = len(filtered_questions).
    if available_questions == 0:
        return []
    if available_questions < num_questions:
        num_questions = available_questions
    sampled_questions = filtered_questions.sample(n=num_questions, replace=False)
    return [
        {"question": row["question"], "marks": marks}
        for _, row in sampled_questions.iterrows()
    ]



def export_to_word_randomized(selected_sections, exam_type, year, subject, instructions, logo_path=None, num_questions_per_section=3):
    """
    Generate a randomized question paper in a Word document.
    
    Parameters:
        selected_sections (dict): Sections with their respective question lists.
        exam_type (str): Exam type (e.g., Midterm, Final).
        year (str): Year of the exam.
        subject (str): Subject of the exam.
        instructions (list): List of instructions for the exam.
        logo_path (str, optional): Path to the logo image.
        num_questions_per_section (int, optional): Number of questions to select from each section.
    
    Returns:
        BytesIO: The generated Word document.
    """
    doc = Document()

    # Add logo if provided
    if logo_path:
        header = doc.sections[0].header
        logo_paragraph = header.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(1.5))
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add title and exam details
    doc.add_heading("DR. BABASAHEB AMBEDKAR TECHNOLOGICAL UNIVERSITY, LONERE", level=1)
    details_paragraph =doc.add_paragraph(f"{exam_type.upper()} EXAMINATION - {year} Year")
    details_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f"Course: B.Tech.   Branch: Computer Science & Engineering      Semester:", level=2)
    doc.add_heading(f"Subject Code & Name: {subject.upper()}", level=2)
    doc.add_paragraph("Time:  hours\t\tFull Marks: ", style="Normal")

    
        # Add instruction table
    doc.add_heading("Instructions:", level=2)
    instruction_table = doc.add_table(rows=len(instructions), cols=1)
    instruction_table.style = 'Table Grid'

    for i, instruction in enumerate(instructions, start=1):
        instruction_row = instruction_table.rows[i - 1].cells
        instruction_row[0].text = f"{i}. {instruction}"
        for paragraph in instruction_row[0].paragraphs:
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add questions in a single table
    doc.add_heading("Questions:", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

     # Center align the table on the page
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


    # Set table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Q.No."
    hdr_cells[1].text = "Question"
    hdr_cells[2].text = "CO"  # Placeholder for Course Outcome
    hdr_cells[3].text = "Mark"

    table.columns[0].width = Inches(0.5)  # Question No.
    table.columns[1].width = Inches(4.5)  # Question
    table.columns[2].width = Inches(0.5)  # Marks
    table.columns[3].width = Inches(0.5) 

    # Set fixed column widths
    for row in table.rows:
        row.cells[0].width = Cm(2.0)  # Question number
        row.cells[1].width = Cm(12.0)  # Question text
        row.cells[2].width = Cm(2.0)  # Course outcome
        row.cells[3].width = Cm(2.0)  # Mark



# Helper function to format cell
    def format_cell(cell, space_before=6, space_after=6):
        # Center align vertically
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

        # Set spacing and horizontal alignment
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(space_before)
            paragraph.paragraph_format.space_after = Pt(space_after)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add randomized questions
    question_no = 1
    for section, questions in selected_sections.items():
        # Shuffle and select a subset of questions
        random.shuffle(questions)
        selected_questions = questions[:num_questions_per_section]



 # Add all questions divided by main questions
    question_no = 1  # Main question number
    for section, questions in selected_sections.items():
        # Add main question row
        main_question_row = table.add_row().cells
        main_question_row[0].text = f"Q {question_no}"  # Main question number
        main_question_row[1].text = f"Solve any Two"  # Main question title
        main_question_row[1].paragraphs[0].runs[0].bold = True  # Make it bold
        main_question_row[2].text = ""  # Empty CO cell for main question
        main_question_row[3].text = "12"  # Empty Marks cell for main question
        
# Format main question row cells
        for cell in main_question_row:
            format_cell(cell, space_before=12, space_after=12)


        # Add sub-questions under the main question
        for sub_index, q in enumerate(questions, start=1):
            sub_question_row = table.add_row().cells
            sub_question_row[0].text = f"{chr(96 + sub_index).upper()}"  # Format: 1.A, 1.B, etc.
            sub_question_row[1].text = q["question"]
            sub_question_row[2].text = "1"  # Example CO; customize as needed
            sub_question_row[3].text = str(q["marks"])

           # Format sub-question row cells
            for cell in sub_question_row:
                format_cell(cell)

        # Increment the main question number after processing the current section
        question_no += 1

    # Save the document to a buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.title("Question Paper Generator")

# Sidebar: Select the type of question paper
st.sidebar.header("Choose an Option")
exam_type = st.sidebar.radio(
    "Generate Question Paper:",
    options=["Midsem", "Unit Test", "Semester"]
)

# st.sidebar.header("MCQ Settings")
# num_mcqs = st.sidebar.slider("Number of MCQs", 1, 10, 5)
# mcq_marks = st.sidebar.slider("Marks per MCQ", 1, 5, 2)


# Dynamic year and subject selection
st.sidebar.header("Filters")
year = st.sidebar.selectbox("Select Year", options=questions_df['year'].unique())

# Filter subjects based on selected year
filtered_subjects = questions_df[questions_df['year'] == year]['subject'].unique()
subject = st.sidebar.selectbox("Select Subject", options=filtered_subjects)

# Difficulty levels and units
difficulty_levels = st.sidebar.multiselect("Select Difficulty Levels", options=["Easy", "Medium", "Hard"])
selected_units = st.sidebar.multiselect("Select Units", options=["Unit 1", "Unit 2", "Unit 3", "Unit 4", "Unit 5"])

# Instructions input
instructions = st.sidebar.text_area(
    "Enter main instructions (one per line):",
    "All questions are compulsory.\nWrite your answers clearly.\nEach question carries specified marks."
).split("\n")

# Map units to sections
unit_to_section_map = {
    "Unit 1": "Section A",
    "Unit 2": "Section B",
    "Unit 3": "Section C",
    "Unit 4": "Section D",
    "Unit 5": "Section E",
}

# Number of questions and marks sliders for each section
section_settings = {}
for unit in selected_units:
    st.sidebar.header(f"Settings for {unit}")
    num_questions = st.sidebar.slider(f"Questions in {unit}", 0, 5, 3)
    marks_per_question = st.sidebar.slider(f"Marks per Question in {unit}", 1, 10, 5)
    section_settings[unit] = {
        "num_questions": num_questions,
        "marks": marks_per_question,
    }

# Function to generate questions
def generate_question_paper():
    selected_sections_questions = {}
    mcq_questions = []
    for unit, settings in section_settings.items():
        # Filter questions based on subject and difficulty level
        unit_questions = questions_df[(questions_df['subject'] == subject) & 
                                      (questions_df['units'] == unit) & 
                                      (questions_df['difficulty'].isin(difficulty_levels))]
        questions = generate_questions(unit_questions, settings["num_questions"], settings["marks"])
        selected_sections_questions[unit_to_section_map[unit]] = questions
         
    # Display in Streamlit
    st.write(f"## {exam_type} Question Paper")
    for i, (section, questions) in enumerate(selected_sections_questions.items(), start=1):
        st.write(f"### Question {i}:")
        for j, q in enumerate(questions, start=1):
            st.write(f"{chr(96 + j).upper()}. {q['question']} [{q['marks']} Marks]")




    # Export to Word
    buffer = export_to_word_randomized(selected_sections_questions, exam_type, year, subject, instructions)
    st.download_button(
        label=f"Download {exam_type} Question Paper",
        data=buffer,
        file_name=f"{year}_{subject}_{exam_type}_question_paper.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Generate the selected question paper
if st.sidebar.button(f"Generate {exam_type} Question Paper"):
    generate_question_paper()
